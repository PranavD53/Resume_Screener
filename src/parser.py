import re
import os
import ast
import pickle
from datetime import datetime

# Helper to load skills vocabulary
def load_skills_vocab(vocab_path='models/skills_vocab.pkl'):
    if os.path.exists(vocab_path):
        with open(vocab_path, 'rb') as f:
            return pickle.load(f)
    return []

# Extract text from PDF using pypdf
def extract_text_from_pdf(file_path):
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        print(f"Error reading PDF '{file_path}': {e}")
        return ""

# Extract text from plain text files
def extract_text_from_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        print(f"Error reading TXT '{file_path}': {e}")
        return ""

# Extract text from DOCX (paragraphs + tables)
def extract_text_from_docx(file_path):
    try:
        import docx
        doc = docx.Document(file_path)
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                
        # Extract from tables (often used for education/experience layouts)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_val = cell.text.strip()
                    if cell_val and cell_val not in row_text: # avoid duplicate text in merged cells
                        row_text.append(cell_val)
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
        return "\n".join(text_parts)
    except Exception as e:
        print(f"Error reading docx '{file_path}': {e}")
        return ""

# Extract Email
def extract_email(text):
    try:
        email_regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        match = re.search(email_regex, text)
        if match:
            return match.group(0).strip()
    except Exception as e:
        print(f"Email extraction error: {e}")
    return "N/A"

# Extract Phone
def extract_phone(text):
    try:
        phone_patterns = [
            r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', # US format
            r'\b\+?\d{1,4}[-.\s]?\d{9,10}\b', # Continuous international
            r'\b\+?\d{1,3}[-.\s]?\d{4,5}[-.\s]?\d{4,5}\b' # Split international
        ]
        for pattern in phone_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0).strip()
    except Exception as e:
        print(f"Phone extraction error: {e}")
    return "N/A"

# Extract Location
def extract_location(text):
    try:
        location_patterns = [
            r'(?:location|address|residence|reside\s+in|lives\s+in)\s*:\s*([^\n,|]+(?:,\s*[^\n,|]+)?)',
            r'\b([A-Z][a-zA-Z\s]+,\s*[A-Z]{2}(?:\s+\d{5})?)\b' # E.g. Woonsocket, RI
        ]
        for pattern in location_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                loc = match.group(1).strip()
                if len(loc) < 50 and not any(x in loc.lower() for x in ['email', 'phone', 'resume', 'skills']):
                    return loc
                    
        # Fallback first city-state line
        lines = [line.strip() for line in text.split('\n') if line.strip()][:15]
        for line in lines:
            if re.search(r'\b[A-Z][a-z]+,\s*[A-Z][a-z]+\b|\b[A-Z][a-z]+,\s*[A-Z]{2}\b', line):
                return line
    except Exception as e:
        print(f"Location extraction error: {e}")
    return "N/A"

# Extract Candidate Name
def extract_name(text, file_path=None):
    try:
        text_clean = text.replace('\r', '\n')
        lines = [line.strip() for line in text_clean.split('\n') if line.strip()][:10]
        
        for line in lines:
            cleaned_line = re.sub(r'^(?:name|candidate\s*name|candidate)\s*:\s*', '', line, flags=re.IGNORECASE).strip()
            if not cleaned_line or len(cleaned_line) > 40:
                continue
                
            if any(x in cleaned_line.lower() for x in ['email', 'phone', 'mobile', 'address', 'resume', 'cv', 'curriculum', 'page', 'profile', 'objective', 'link', 'github', 'linkedin', 'http', '@', 'location', 'nationality', 'born']):
                continue
                
            # Clean punctuation from words
            words = [re.sub(r'[^\w]', '', w) for w in cleaned_line.split()]
            words = [w for w in words if w]
            
            if 2 <= len(words) <= 4:
                # Check capitalization of alphabetic words
                cap_words = [w for w in words if w[0].isupper() if w[0].isalpha()]
                if len(cap_words) >= 2:
                    name_candidate = cleaned_line.strip(',. |')
                    if not any(char.isdigit() for char in name_candidate):
                        return name_candidate
    except Exception as e:
        print(f"Name extraction error: {e}")
        
    # Fallback to filename
    if file_path:
        try:
            base = os.path.basename(file_path)
            name_part = os.path.splitext(base)[0]
            name_clean = name_part.replace('_', ' ').replace('-', ' ')
            name_clean = re.sub(r'\b(?:resume|cv|profile|java|fullstack|developer|qa|pm|sm|bsa|agile|project|manager|updated|latest|new|draft|copy)\b', '', name_clean, flags=re.IGNORECASE)
            name_clean = re.sub(r'\s+', ' ', name_clean).strip()
            words = [w.capitalize() for w in name_clean.split() if w]
            if len(words) >= 2:
                return " ".join(words[:3])
        except:
            pass
            
    return "Candidate"

# Parsing Education Level helper
def parse_education_level(edu_str):
    if not isinstance(edu_str, str) or not edu_str.strip():
        return 1
    edu_str = edu_str.lower()
    # Word boundaries to prevent matching fragments (e.g. amber -> mba)
    if re.search(r'\b(?:phd|doctorate|doctor\s+of\s+philosophy|ph\.d)\b', edu_str):
        return 4
    if re.search(r'\b(?:masters?|m\.s|m\.sc|msc|mtech|m\.tech|mba|post\s*graduate|p\.g\b|pg\b)\b', edu_str):
        return 3
    if re.search(r'\b(?:bachelors?|b\.s|b\.sc|bsc|btech|b\.tech|b\.a|be|b\.e|undergraduate|u\.g\b|ug\b)\b', edu_str):
        return 2
    return 1

# Extract Education
def extract_education(text):
    try:
        level = parse_education_level(text)
        levels_map = {
            4: "Doctorate (PhD)",
            3: "Master's Degree",
            2: "Bachelor's Degree",
            1: "High School / Other"
        }
        return levels_map.get(level, "High School / Other"), level
    except Exception as e:
        print(f"Education extraction error: {e}")
    return "High School / Other", 1

# Parse Required Experience helper
def parse_required_experience(req_str):
    if not isinstance(req_str, str):
        return 0.0
    req_str = req_str.lower()
    range_match = re.search(r'(\d+)\s*(?:to|-)\s*(\d+)', req_str)
    if range_match:
        return float(range_match.group(1))
    match = re.search(r'(\d+)\s*(?:years?|yrs?|year)', req_str)
    if match:
        return float(match.group(1))
    return 0.0

# Extract Experience in Years (combines text mentions and overlapping date intervals)
def extract_experience(text):
    try:
        # Normalize text dashes and spaces for date matching
        text_clean = text.lower()
        text_clean = text_clean.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-').replace('–', '-').replace('—', '-')
        text_clean = text_clean.replace('\u00a0', ' ').replace('&nbsp;', ' ')
        text_clean = re.sub(r'\s+', ' ', text_clean)
        
        # Regex to capture months and years
        month_numeric = r'(?:\d{1,2}[-/\s]+)'
        month_text = r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*,?\s*'
        date_pat = rf'(?:{month_text}|{month_numeric})?\b(?:19\d{{2}}|20\d{{2}})\b'
        present_pat = r'(?:present|current|till\s*date|now|ongoing|till\s*now)'
        
        span_regex = rf'({date_pat})\s*(?:-|to|until)\s*({date_pat}|{present_pat})'
        spans = re.findall(span_regex, text_clean)
        
        def parse_dt(d_str):
            d_str = d_str.strip()
            if re.match(rf'^{present_pat}$', d_str):
                return datetime.now()
            # Standard formatting cleanups
            for fmt in ('%b %Y', '%B %Y', '%m/%Y', '%Y'):
                try:
                    clean_str = re.sub(r'\s+', ' ', d_str)
                    return datetime.strptime(clean_str, fmt)
                except:
                    pass
            # Try numeric formats like 11-2019
            dash_match = re.search(r'\b(\d{1,2})[-/\s]+(19\d{2}|20\d{2})\b', d_str)
            if dash_match:
                return datetime(int(dash_match.group(2)), int(dash_match.group(1)), 1)
            yr = re.search(r'\b(19\d{2}|20\d{2})\b', d_str)
            if yr:
                return datetime(int(yr.group(1)), 1, 1)
            return None

        intervals = []
        for start, end in spans:
            s_dt = parse_dt(start)
            e_dt = parse_dt(end)
            if s_dt and e_dt and s_dt <= e_dt:
                intervals.append((s_dt, e_dt))
                
        # Merge overlapping intervals to prevent double-counting concurrent roles
        calc_exp = 0.0
        if intervals:
            intervals.sort(key=lambda x: x[0])
            merged = [intervals[0]]
            for current in intervals[1:]:
                prev = merged[-1]
                if current[0] <= prev[1]:
                    merged[-1] = (prev[0], max(prev[1], current[1]))
                else:
                    merged.append(current)
            
            total_months = 0
            for start, end in merged:
                months = (end.year - start.year) * 12 + (end.month - start.month)
                if 0 < months < 600:
                    total_months += months
            calc_exp = round(total_months / 12.0, 1)
            
        # Scan for direct text mentions like "X+ years of experience"
        mention_regex = r'\b(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|work|industry|professional)'
        mentions = re.findall(mention_regex, text_clean)
        mention_exp = 0.0
        if mentions:
            try:
                mention_exp = max([float(m) for m in mentions])
            except:
                pass
                
        return min(max(calc_exp, mention_exp), 40.0)
    except Exception as e:
        print(f"Experience extraction error: {e}")
    return 0.0

# Helper to check boundary for skills including symbols (e.g. C++, C#)
def get_skill_pattern(skill):
    skill_lower = skill.lower()
    # Ends with non-word character (like +, #)
    if not skill_lower[-1].isalnum():
        return rf'\b{re.escape(skill_lower)}(?!\w)'
    # Starts with non-word character (like .)
    elif not skill_lower[0].isalnum():
        return rf'(?<!\w){re.escape(skill_lower)}\b'
    else:
        return rf'\b{re.escape(skill_lower)}\b'

# Extract Skills using the vocabulary
def extract_skills(text, skills_vocab):
    if not skills_vocab:
        return []
    try:
        # Normalize all whitespaces (including newlines and multiple spaces)
        # to a single space, preventing split words in PDFs or TXT files
        text_clean = re.sub(r'\s+', ' ', text.lower())
        text_lower = text_clean
        extracted = []
        
        for skill in skills_vocab:
            pattern = get_skill_pattern(skill)
            if re.search(pattern, text_lower):
                extracted.append(skill)
            elif len(skill) > 5 and skill.lower() in text_lower:
                # Fallback for longer skill phrases
                extracted.append(skill)
                
        return sorted(list(set(extracted)))
    except Exception as e:
        print(f"Skills extraction error: {e}")
    return []

# Expose skill extractor for import
def extract_skills_from_text(text, skills_vocab):
    return extract_skills(text, skills_vocab)

# Main parse interface
def parse_resume(file_path, skills_vocab=None):
    if skills_vocab is None:
        skills_vocab = load_skills_vocab()
        
    # Check extension
    ext = file_path.lower()
    if ext.endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif ext.endswith('.txt'):
        text = extract_text_from_txt(file_path)
    else:
        text = extract_text_from_docx(file_path)
        
    name = extract_name(text, file_path)
    email = extract_email(text)
    phone = extract_phone(text)
    location = extract_location(text)
    
    edu_name, edu_level = extract_education(text)
    exp_years = extract_experience(text)
    skills = extract_skills(text, skills_vocab)
    
    return {
        'filename': os.path.basename(file_path),
        'filepath': file_path,
        'name': name,
        'email': email,
        'phone': phone,
        'location': location,
        'education_degree': edu_name,
        'education_level': edu_level,
        'experience_years': exp_years,
        'skills': skills,
        'text': text
    }
